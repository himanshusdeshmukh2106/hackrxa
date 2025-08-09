"""
Text extraction service for different document formats
"""
import io
import email
import re
import asyncio
import os
import time
from typing import List, Dict, Any, Optional
from email.message import EmailMessage

import PyPDF2
import pdfplumber
import fitz  # PyMuPDF - much faster
from concurrent.futures import ThreadPoolExecutor
from docx import Document as DocxDocument

from app.core.config import settings
from app.core.logging import LoggerMixin
from app.core.exceptions import DocumentProcessingError
from app.schemas.models import Document, TextChunk, DocumentType
from app.core.cache import cache_manager


class TextExtractor(LoggerMixin):
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.max_chunk_size = settings.max_chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.pdf_batch_size = settings.pdf_batch_size
    
    async def extract_text(self, document: Document) -> List[TextChunk]:
        """
        Extract text from document and return chunks
        
        Args:
            document: Document object with raw content
            
        Returns:
            List of TextChunk objects
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            raw_content = document.metadata.get("raw_content")
            if not raw_content:
                raise DocumentProcessingError("No raw content found in document")
            
            # Extract text based on document type
            if document.document_type == DocumentType.PDF:
                text = await self._extract_pdf_text(raw_content)
            elif document.document_type == DocumentType.DOCX:
                text = await self._extract_docx_text(raw_content)
            elif document.document_type == DocumentType.EMAIL:
                text = await self._extract_email_text(raw_content)
            elif document.document_type == DocumentType.TEXT:
                text = await self._extract_plain_text(raw_content)
            else:
                raise DocumentProcessingError(f"Unsupported document type: {document.document_type}")
            
            if not text.strip():
                raise DocumentProcessingError("No text content extracted from document")
            
            # Create text chunks
            chunks = self._create_text_chunks(document.id, text)
            
            self.logger.info(f"Extracted {len(chunks)} text chunks from document {document.id}")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Text extraction failed for document {document.id}: {str(e)}")
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")
    
    async def extract_and_chunk_text_fast(self, document: Document) -> List[TextChunk]:
        """
        Fast text extraction and chunking with optimizations
        """
        try:
            # Extract text with timeout
            text_chunks = await asyncio.wait_for(
                self.extract_text(document), 
                timeout=15.0  # Reduced timeout for speed
            )
            
            # Return chunks directly (already chunked by extract_text)
            return text_chunks
            
        except asyncio.TimeoutError:
            self.logger.error(f"Fast text extraction timed out for document {document.id}")
            raise DocumentProcessingError("Text extraction timed out")
        except Exception as e:
            self.logger.error(f"Fast text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"Fast extraction failed: {str(e)}")
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """
        Ultra-fast PDF text extraction using PyMuPDF with parallel processing
        
        Args:
            content: Raw PDF bytes
            
        Returns:
            Extracted text string
        """
        start_time = time.monotonic()
        try:
            # Use PyMuPDF (fitz) - 3-5x faster than pdfplumber
            doc = fitz.open(stream=content, filetype="pdf")
            page_count = doc.page_count
            
            self.logger.info(f"Fast PDF extraction: {page_count} pages using PyMuPDF")
            
            if page_count <= 10:
                # Small PDFs - extract directly (fastest path)
                text_parts = []
                for page_num in range(page_count):
                    page = doc.load_page(page_num)
                    page_text = page.get_text("text")  # Text-only mode (fastest)
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}] {page_text}")
                
                doc.close()
                return "\n\n".join(text_parts)
            
            else:
                # Large PDFs - parallel processing in batches
                batch_size = self.pdf_batch_size  # Process pages per batch
                batches = [range(i, min(i + batch_size, page_count)) 
                          for i in range(0, page_count, batch_size)]
                
                # Use thread pool for parallel extraction
                with ThreadPoolExecutor(max_workers=min(64, os.cpu_count() * 2 + 4)) as executor:
                    batch_results = list(executor.map(
                        lambda batch: self._extract_page_batch(content, batch), 
                        batches
                    ))
                
                doc.close()
                
                # Combine results
                all_text = []
                for batch_text in batch_results:
                    if batch_text:
                        all_text.append(batch_text)
                
                return "\n\n".join(all_text)
        
        except Exception as e:
            self.logger.error(f"PyMuPDF extraction failed: {str(e)}")
            
            # Fallback to pdfplumber for complex layouts
            try:
                self.logger.info("Falling back to pdfplumber")
                fallback_start_time = time.monotonic()
                result = await self._extract_pdf_fallback(content)
                fallback_duration = time.monotonic() - fallback_start_time
                self.logger.info(f"pdfplumber fallback took {fallback_duration:.2f} seconds")
                return result
            except Exception as fallback_error:
                self.logger.error(f"Fallback extraction failed: {str(fallback_error)}")
                raise DocumentProcessingError(f"PDF extraction failed: {str(e)}")
        finally:
            duration = time.monotonic() - start_time
            self.logger.info(f"PyMuPDF extraction finished in {duration:.2f} seconds")
    
    def _extract_page_batch(self, content: bytes, page_range: range) -> str:
        """
        Extract text from a batch of pages (for parallel processing)
        
        Args:
            content: PDF content bytes
            page_range: Range of page numbers to extract
            
        Returns:
            Combined text from the page batch
        """
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            
            for page_num in page_range:
                page = doc.load_page(page_num)
                page_text = page.get_text("text")  # Fast text-only extraction
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}] {page_text}")
            
            doc.close()
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Batch extraction failed for pages {page_range}: {str(e)}")
            return ""
    
    async def _extract_pdf_fallback(self, content: bytes) -> str:
        """
        Fallback PDF extraction using pdfplumber (for complex layouts)
        """
        text_parts = []
        
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}] {page_text}")
            
            if text_parts:
                return "\n\n".join(text_parts)
        
        except Exception as e:
            self.logger.warning(f"pdfplumber fallback failed: {str(e)}")
        
        # Final fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}] {page_text}")
            
            if text_parts:
                return "\n\n".join(text_parts)
        
        except Exception as e:
            self.logger.error(f"All PDF extraction methods failed: {str(e)}")
        
        raise DocumentProcessingError("No text could be extracted from PDF")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """
        Extract text from DOCX content
        
        Args:
            content: Raw DOCX bytes
            
        Returns:
            Extracted text string
        """
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_text = f"\n[Table {table_num + 1}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text += row_text + "\n"
                text_parts.append(table_text)
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text_parts.insert(0, f"[Header] {para.text}")
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[Footer] {para.text}")
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            raise DocumentProcessingError(f"DOCX text extraction failed: {str(e)}")
    
    async def _extract_email_text(self, content: bytes) -> str:
        """
        Extract text from email content
        
        Args:
            content: Raw email bytes
            
        Returns:
            Extracted text string
        """
        try:
            # Parse email
            email_str = content.decode('utf-8', errors='ignore')
            msg = email.message_from_string(email_str)
            
            text_parts = []
            
            # Extract headers
            headers = [
                f"From: {msg.get('From', 'Unknown')}",
                f"To: {msg.get('To', 'Unknown')}",
                f"Subject: {msg.get('Subject', 'No Subject')}",
                f"Date: {msg.get('Date', 'Unknown')}",
            ]
            
            # Add CC and BCC if present
            if msg.get('Cc'):
                headers.append(f"Cc: {msg.get('Cc')}")
            if msg.get('Bcc'):
                headers.append(f"Bcc: {msg.get('Bcc')}")
            
            text_parts.append("[Email Headers]\n" + "\n".join(headers) + "\n")
            
            # Extract body
            body_text = self._extract_email_body(msg)
            if body_text:
                text_parts.append("[Email Body]\n" + body_text)
            
            # Extract attachment info
            attachments = self._extract_attachment_info(msg)
            if attachments:
                text_parts.append("[Attachments]\n" + "\n".join(attachments))
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            raise DocumentProcessingError(f"Email text extraction failed: {str(e)}")
    
    def _extract_email_body(self, msg: EmailMessage) -> str:
        """Extract body text from email message"""
        body_parts = []
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_parts.append(payload.decode('utf-8', errors='ignore'))
                elif content_type == "text/html":
                    # Basic HTML to text conversion
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_text = payload.decode('utf-8', errors='ignore')
                        plain_text = self._html_to_text(html_text)
                        body_parts.append(f"[HTML Content]\n{plain_text}")
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body_parts.append(payload.decode('utf-8', errors='ignore'))
        
        return "\n\n".join(body_parts)
    
    def _extract_attachment_info(self, msg: EmailMessage) -> List[str]:
        """Extract attachment information from email"""
        attachments = []
        
        for part in msg.walk():
            if part.get_content_disposition() == 'attachment':
                filename = part.get_filename()
                if filename:
                    content_type = part.get_content_type()
                    size = len(part.get_payload(decode=True) or b'')
                    attachments.append(f"- {filename} ({content_type}, {size} bytes)")
        
        return attachments
    
    def _html_to_text(self, html: str) -> str:
        """Basic HTML to text conversion"""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', html)
        # Decode HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&amp;', '&')
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    async def _extract_plain_text(self, raw_content: str) -> str:
        """Extract text from plain text file"""
        return raw_content
    
    def _format_table_text(self, table: List[List[str]], page_num: int, table_num: int) -> str:
        """Format table data as text"""
        if not table:
            return ""
        
        text_parts = [f"\n[Page {page_num} - Table {table_num}]"]
        
        for row in table:
            if row:
                # Clean and join cells
                clean_row = [str(cell).strip() if cell else "" for cell in row]
                if any(clean_row):  # Only add non-empty rows
                    text_parts.append(" | ".join(clean_row))
        
        return "\n".join(text_parts) + "\n"
    
    def _create_text_chunks(self, document_id: str, text: str) -> List[TextChunk]:
        """
        Split text into chunks with overlap
        
        Args:
            document_id: Document ID
            text: Full text to chunk
            
        Returns:
            List of TextChunk objects
        """
        chunks = []
        
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= self.max_chunk_size:
            # Single chunk
            chunk = TextChunk(
                document_id=document_id,
                content=text,
                chunk_index=0,
                start_char=0,
                end_char=len(text),
                metadata={"chunk_method": "single"}
            )
            chunks.append(chunk)
        else:
            # Multiple chunks with overlap
            start = 0
            chunk_index = 0
            
            while start < len(text):
                end = min(start + self.max_chunk_size, len(text))
                
                # Try to break at sentence boundary
                if end < len(text):
                    # Look for sentence endings within the last 200 characters
                    search_start = max(end - 200, start)
                    sentence_end = max(
                        text.rfind('.', search_start, end),
                        text.rfind('!', search_start, end),
                        text.rfind('?', search_start, end)
                    )
                    
                    if sentence_end > start:
                        end = sentence_end + 1
                
                chunk_text = text[start:end].strip()
                
                if chunk_text:
                    chunk = TextChunk(
                        document_id=document_id,
                        content=chunk_text,
                        chunk_index=chunk_index,
                        start_char=start,
                        end_char=end,
                        metadata={
                            "chunk_method": "overlap",
                            "overlap_size": self.chunk_overlap
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Move start position with overlap
                start = max(start + self.max_chunk_size - self.chunk_overlap, end)
        
        return chunks